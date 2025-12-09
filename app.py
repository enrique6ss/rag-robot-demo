import os
import io
import numpy as np
import streamlit as st

from llama_index.core import VectorStoreIndex, Settings, Document
from llama_index.llms.groq import Groq
from llama_index.embeddings.openai import OpenAIEmbedding

import pdfplumber
from pdf2image import convert_from_path
import easyocr
from docx import Document as DocxDocument

# -----------------------------
# Page + Global Configuration
# -----------------------------

st.set_page_config(
    page_title="LexiScan AI – Document Intelligence Platform",
    layout="wide",
)

# Custom dark blue gradient UI
st.markdown(
    """
    <style>
    body {
        background: radial-gradient(circle at top left, #1e293b 0, #020617 55%, #020617 100%) !important;
        color: #e5e7eb !important;
    }
    .block-container {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
        max-width: 1100px !important;
    }
    .lexi-card {
        background: rgba(15,23,42,0.96);
        border-radius: 20px;
        padding: 1.5rem 2rem 1.75rem 2rem;
        border: 1px solid rgba(148,163,184,0.55);
        box-shadow: 0 28px 80px rgba(15,23,42,0.95);
        backdrop-filter: blur(22px);
    }
    .lexi-pill {
        display:inline-flex;
        align-items:center;
        gap:0.5rem;
        padding:0.2rem 0.9rem;
        border-radius:999px;
        background:linear-gradient(120deg,#0ea5e9,#22c55e);
        font-size:0.72rem;
        font-weight:700;
        text-transform:uppercase;
        letter-spacing:0.16em;
        color:#020617;
    }
    .lexi-pill span {
        font-size:0.72rem;
    }
    h1.lexi-title {
        margin-top:0.85rem;
        font-size:2.3rem;
        line-height:1.1;
        letter-spacing:-0.03em;
        color:#e5e7eb;
    }
    p.lexi-subtitle {
        color:#c7d2fe;
        max-width:640px;
        margin-top:0.6rem;
        font-size:0.95rem;
    }
    .lexi-metrics {
        display:flex;
        flex-wrap:wrap;
        gap:1.25rem;
        margin-top:1.15rem;
        font-size:0.8rem;
        color:#9ca3af;
    }
    .lexi-metric-pill {
        padding:0.4rem 0.8rem;
        border-radius:999px;
        border:1px solid rgba(148,163,184,0.5);
        background:rgba(15,23,42,0.8);
    }
    .stChatMessage {
        border-radius: 16px;
        border: 1px solid rgba(148,163,184,0.25);
        background: rgba(15,23,42,0.92) !important;
    }
    .stChatMessage[data-testid="stChatMessageUser"] {
        border-color: rgba(56,189,248,0.6);
    }
    .stChatMessage[data-testid="stChatMessageAssistant"] {
        border-color: rgba(129,140,248,0.7);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# -----------------------------
# Model / Embedding Setup
# -----------------------------

llm = Groq(
    model="llama-3.3-70b-versatile",
    api_key=os.getenv("GROQ_API_KEY"),
)
Settings.llm = llm
Settings.embed_model = OpenAIEmbedding(api_key=os.getenv("OPENAI_API_KEY"))

DATA_FOLDER = "data"
os.makedirs(DATA_FOLDER, exist_ok=True)

# -----------------------------
# OCR + Text Extraction Helpers
# -----------------------------

@st.cache_resource
def get_easyocr_reader():
    # English only for speed; set gpu=False for Railway
    return easyocr.Reader(["en"], gpu=False)


def extract_text_from_pdf(filepath: str) -> str:
    """
    Hybrid PDF text extraction:
    1) Try native text via pdfplumber.
    2) If almost empty, fallback to EasyOCR over rendered images.
    """
    text_chunks = []

    # Step 1: native text
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_chunks.append(page_text)
    except Exception:
        # If any problem, we'll rely on OCR fallback
        pass

    native_text = "\n".join(text_chunks).strip()

    # If we got enough real text, use it
    if len(native_text) > 80:
        return native_text

    # Step 2: EasyOCR fallback
    reader = get_easyocr_reader()
    ocr_text_parts = []
    try:
        images = convert_from_path(filepath)
        for img in images:
            # EasyOCR expects file path, bytes, or numpy array
            img_np = np.array(img)
            result = reader.readtext(img_np, detail=0)
            if result:
                ocr_text_parts.append(" ".join(result))
    except Exception:
        pass

    ocr_text = "\n".join(ocr_text_parts).strip()

    # Prefer OCR text if we had almost no native text
    if ocr_text:
        return ocr_text

    # Fallback if nothing at all
    return native_text or "[No text detected in this PDF]"


def extract_text_from_docx(filepath: str) -> str:
    try:
        doc = DocxDocument(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def extract_text_from_txt(filepath: str) -> str:
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def build_index_from_files() -> VectorStoreIndex:
    """
    Build a VectorStoreIndex from all files in DATA_FOLDER
    using hybrid text + OCR extraction.
    """
    docs = []

    for filename in sorted(os.listdir(DATA_FOLDER)):
        filepath = os.path.join(DATA_FOLDER, filename)
        if not os.path.isfile(filepath):
            continue
        if filename.startswith("."):
            continue

        ext = filename.lower().split(".")[-1]

        if ext == "pdf":
            text = extract_text_from_pdf(filepath)
        elif ext == "txt":
            text = extract_text_from_txt(filepath)
        elif ext == "docx":
            text = extract_text_from_docx(filepath)
        else:
            continue  # ignore unsupported types

        if not text or not text.strip():
            continue

        # New Document API: use keyword args (works across versions)
        docs.append(
            Document(
                text=text,
                metadata={"filename": filename},
            )
        )

    if not docs:
        # Safe dummy document; no weird positional args
        docs.append(
            Document(
                text="Upload a document to begin.",
                metadata={"filename": "none"},
            )
        )

    return VectorStoreIndex.from_documents(docs)

# -----------------------------
# Header / Hero Section
# -----------------------------

st.markdown(
    """
    <div class="lexi-card">
        <div class="lexi-pill">
            <span>LexiScan AI</span> · <span>Document Intelligence</span>
        </div>
        <h1 class="lexi-title">LexiScan AI – Document Intelligence Platform</h1>
        <p class="lexi-subtitle">
            Upload contracts, NDAs, and legal documents. Ask questions in plain English and
            get precise, citation-backed answers in seconds – powered by hybrid OCR + text search.
        </p>
        <div class="lexi-metrics">
            <div class="lexi-metric-pill">⚖️ Built for law & real estate workflows</div>
            <div class="lexi-metric-pill">🧠 Groq 70B + OpenAI embeddings</div>
            <div class="lexi-metric-pill">📄 PDF / DOCX / TXT + OCR fallback</div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.markdown("")  # spacing

# -----------------------------
# Sidebar: Controls
# -----------------------------

st.sidebar.header("Workspace")

top_k = st.sidebar.slider("Relevant chunks per answer", 1, 8, 3)

if st.sidebar.button("Clear chat history"):
    st.session_state.pop("messages", None)

# -----------------------------
# File Upload + Index Rebuild
# -----------------------------

uploaded_files = st.file_uploader(
    "Upload documents (PDF / DOCX / TXT)",
    accept_multiple_files=True,
    type=["pdf", "docx", "txt"],
)

if uploaded_files:
    for f in uploaded_files:
        save_path = os.path.join(DATA_FOLDER, f.name)
        with open(save_path, "wb") as out:
            out.write(f.getbuffer())

    st.success("✅ Files uploaded. Rebuilding intelligence index...")
    with st.spinner("Rebuilding LexiScan intelligence index…"):
        st.session_state["index"] = build_index_from_files()

# On first run, ensure we have an index
if "index" not in st.session_state:
    with st.spinner("Initializing LexiScan intelligence index…"):
        st.session_state["index"] = build_index_from_files()

index: VectorStoreIndex = st.session_state["index"]
query_engine = index.as_query_engine(similarity_top_k=top_k)

# -----------------------------
# Chat Interface
# -----------------------------

if "messages" not in st.session_state:
    st.session_state["messages"] = []

# Show history
for message in st.session_state["messages"]:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

prompt = st.chat_input("Ask LexiScan anything about your documents…")

if prompt:
    # User message
    st.session_state["messages"].append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Assistant answer
    with st.chat_message("assistant"):
        with st.spinner("Analyzing documents with LexiScan AI…"):
            response = query_engine.query(prompt)
            answer_text = str(response)

            st.markdown(answer_text)

            # Show simple source list if available
            try:
                if hasattr(response, "source_nodes") and response.source_nodes:
                    st.markdown("##### Sources")
                    seen = set()
                    for node in response.source_nodes:
                        meta = getattr(node, "metadata", None) or getattr(
                            getattr(node, "node", None), "metadata", {}
                        )
                        filename = None
                        if isinstance(meta, dict):
                            filename = meta.get("filename")
                        if not filename:
                            filename = "Document fragment"
                        if filename in seen:
                            continue
                        seen.add(filename)
                        st.markdown(f"- `{filename}`")
            except Exception:
                # Don't break UI if source parsing fails
                pass

            st.download_button(
                label="Download answer as .txt",
                data=answer_text,
                file_name="lexiscan_answer.txt",
                mime="text/plain",
            )

            # Save in history
            st.session_state["messages"].append(
                {"role": "assistant", "content": answer_text}
            )
